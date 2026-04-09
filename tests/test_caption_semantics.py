from io import BytesIO

from docx import Document
from lxml import etree

from src.engine.rules.caption_format import (
    _rebuild_para_as_caption_text,
    _rebuild_para_text_full,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _ids(doc: Document, tag: str) -> list[str]:
    return [
        elem.get(_w("id"))
        for elem in doc.element.body.findall(f".//{_w(tag)}")
        if elem.get(_w("id")) is not None
    ]


def _comment_reference_count(doc: Document) -> int:
    return len(doc.element.body.findall(f".//{_w('commentReference')}"))


def test_caption_rebuild_keeps_cross_paragraph_bookmarks_in_normal_path():
    doc = Document()
    para = doc.add_paragraph("图1.1旧标题")
    tail = doc.add_paragraph("next paragraph")

    start = etree.Element(_w("bookmarkStart"))
    start.set(_w("id"), "191")
    start.set(_w("name"), "OLE_LINK107")
    end = etree.Element(_w("bookmarkEnd"))
    end.set(_w("id"), "191")
    para._element.insert(0, start)
    tail._element.insert(0, end)

    _rebuild_para_as_caption_text(
        para,
        prefix="图",
        chapter_num=1,
        seq_num=2,
        sep=" ",
        title="新标题",
        numbering_format="chapter.seq",
    )

    assert _ids(doc, "bookmarkStart").count("191") == 1
    assert _ids(doc, "bookmarkEnd").count("191") == 1


def test_caption_rebuild_keeps_comment_reference_in_normal_path():
    doc = Document()
    para = doc.add_paragraph("图1.1旧标题")

    start = etree.Element(_w("commentRangeStart"))
    start.set(_w("id"), "7")
    end = etree.Element(_w("commentRangeEnd"))
    end.set(_w("id"), "7")
    ref_run = etree.Element(_w("r"))
    ref = etree.SubElement(ref_run, _w("commentReference"))
    ref.set(_w("id"), "7")
    para._element.insert(0, start)
    para._element.append(ref_run)
    para._element.append(end)

    before_refs = _comment_reference_count(doc)
    _rebuild_para_as_caption_text(
        para,
        prefix="图",
        chapter_num=1,
        seq_num=2,
        sep=" ",
        title="新标题",
        numbering_format="chapter.seq",
    )

    assert _comment_reference_count(doc) == before_refs
    tags = [etree.QName(child).localname for child in para._element]
    assert "commentRangeStart" in tags
    assert "commentRangeEnd" in tags


def test_caption_full_rebuild_keeps_cross_paragraph_bookmarks_in_fallback_path():
    doc = Document()
    para = doc.add_paragraph("")
    tail = doc.add_paragraph("after")

    start = etree.Element(_w("bookmarkStart"))
    start.set(_w("id"), "23")
    start.set(_w("name"), "_Toc23")
    end = etree.Element(_w("bookmarkEnd"))
    end.set(_w("id"), "23")
    para._element.insert(0, start)
    tail._element.insert(0, end)

    _rebuild_para_text_full(
        para._element,
        prefix="图",
        chapter_num=1,
        seq_num=3,
        sep=" ",
        title="补建标题",
        rpr_src=None,
        numbering_format="chapter.seq",
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    reloaded = Document(buf)
    assert _ids(reloaded, "bookmarkStart").count("23") == 1
    assert _ids(reloaded, "bookmarkEnd").count("23") == 1
