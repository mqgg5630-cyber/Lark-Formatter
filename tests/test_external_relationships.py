import os
from pathlib import Path

from docx import Document
from lxml import etree

from src.engine.pipeline import Pipeline
from src.engine.rules.section_format import _normalize_related_story_part_font_hints
from src.scene.manager import load_default_scene


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _add_external_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = etree.SubElement(paragraph._element, f"{{{W_NS}}}hyperlink")
    hyperlink.set(f"{{{R_NS}}}id", rel_id)
    run = etree.SubElement(hyperlink, f"{{{W_NS}}}r")
    text_el = etree.SubElement(run, f"{{{W_NS}}}t")
    text_el.text = text


def _make_external_link_doc() -> Document:
    doc = Document()
    doc.add_paragraph("封面")
    doc.add_paragraph("摘要")
    para = doc.add_paragraph("外链: ")
    _add_external_hyperlink(para, "OpenAI", "https://openai.com")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("正文内容")
    return doc


def test_normalize_related_story_part_font_hints_skips_external_hyperlinks():
    doc = _make_external_link_doc()

    changed = _normalize_related_story_part_font_hints(doc)

    assert changed == {}


def test_pipeline_succeeds_for_doc_with_external_hyperlink(tmp_path):
    src = tmp_path / "external_link.docx"
    _make_external_link_doc().save(src)

    cfg = load_default_scene()
    cfg.output.final_docx = True
    cfg.output.compare_docx = False
    cfg.output.report_json = False
    cfg.output.report_markdown = False

    os.environ["DOCX_DISABLE_FIELD_REFRESH"] = "1"
    os.environ["LARK_FORMATTER_DISABLE_OFFICE_FALLBACK"] = "1"

    result = Pipeline(cfg).run(str(src))

    assert result.success is True
    failures = [
        rec
        for rec in result.tracker.records
        if rec.rule_name == "section_format" and rec.failure_reason
    ]
    assert failures == []
