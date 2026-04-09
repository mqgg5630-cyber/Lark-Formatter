from io import BytesIO

from docx import Document
from lxml import etree

from src.utils.ooxml_paragraph import replace_paragraph_payload_with_omml


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def _make_omml_para():
    omath_para = etree.Element(f"{{{M_NS}}}oMathPara")
    omath = etree.SubElement(omath_para, f"{{{M_NS}}}oMath")
    run = etree.SubElement(omath, f"{{{M_NS}}}r")
    text = etree.SubElement(run, f"{{{M_NS}}}t")
    text.text = "x=y+z"
    return omath_para


def _bookmark_ids(doc: Document, tag: str) -> list[str]:
    return [
        elem.get(_w("id"))
        for elem in doc.element.body.findall(f".//{_w(tag)}")
        if elem.get(_w("id")) is not None
    ]


def test_replace_paragraph_payload_with_omml_keeps_cross_paragraph_bookmark_start():
    doc = Document()
    first = doc.add_paragraph("x = y + z")
    second = doc.add_paragraph("Heading after formula")

    start = etree.Element(_w("bookmarkStart"))
    start.set(_w("id"), "191")
    start.set(_w("name"), "OLE_LINK107")
    end = etree.Element(_w("bookmarkEnd"))
    end.set(_w("id"), "191")

    first._element.insert(0, start)
    second._element.insert(0, end)

    result = replace_paragraph_payload_with_omml(first._element, _make_omml_para())

    assert result.applied is True
    assert _bookmark_ids(doc, "bookmarkStart").count("191") == 1
    assert _bookmark_ids(doc, "bookmarkEnd").count("191") == 1


def test_replace_paragraph_payload_with_omml_keeps_bookmark_wrapping_order():
    doc = Document()
    para = doc.add_paragraph("x = y + z")

    start = etree.Element(_w("bookmarkStart"))
    start.set(_w("id"), "9")
    start.set(_w("name"), "_Toc9")
    end = etree.Element(_w("bookmarkEnd"))
    end.set(_w("id"), "9")

    para._element.insert(0, start)
    para._element.append(end)

    result = replace_paragraph_payload_with_omml(para._element, _make_omml_para())

    assert result.applied is True
    tags = [etree.QName(child).localname for child in para._element]
    assert tags == ["bookmarkStart", "oMathPara", "bookmarkEnd"]


def test_replace_paragraph_payload_with_omml_refuses_comment_ranges():
    doc = Document()
    para = doc.add_paragraph("x = y + z")

    start = etree.Element(_w("commentRangeStart"))
    start.set(_w("id"), "7")
    end = etree.Element(_w("commentRangeEnd"))
    end.set(_w("id"), "7")
    para._element.insert(0, start)
    para._element.append(end)

    result = replace_paragraph_payload_with_omml(para._element, _make_omml_para())

    assert result.applied is False
    assert result.reason == "unsupported_paragraph_markup"
    tags = [etree.QName(child).localname for child in para._element]
    assert tags == ["commentRangeStart", "r", "commentRangeEnd"]


def test_rewritten_paragraph_roundtrips_with_balanced_bookmarks():
    doc = Document()
    first = doc.add_paragraph("x = y + z")
    second = doc.add_paragraph("Heading after formula")

    start = etree.Element(_w("bookmarkStart"))
    start.set(_w("id"), "191")
    start.set(_w("name"), "OLE_LINK107")
    end = etree.Element(_w("bookmarkEnd"))
    end.set(_w("id"), "191")
    first._element.insert(0, start)
    second._element.insert(0, end)

    result = replace_paragraph_payload_with_omml(first._element, _make_omml_para())
    assert result.applied is True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    reloaded = Document(buf)
    assert _bookmark_ids(reloaded, "bookmarkStart").count("191") == 1
    assert _bookmark_ids(reloaded, "bookmarkEnd").count("191") == 1
